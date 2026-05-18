"""
Text extraction helpers used at upload time.

PDFs are NOT handled here — the agent reads PDFs directly via its built-in
Read tool, which sends the PDF as a document block to Claude for native
vision-based parsing (images, tables, layout). For DOCX we extract paragraphs
and tables (rendered as markdown tables). MD/TXT are read as-is.
"""

import os
from typing import Optional

try:
    import docx as _docx
    from docx.oxml.ns import qn as _qn
except ImportError:
    _docx = None
    _qn = None


def _docx_table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ").replace("|", "\\|") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if not rows:
        return ""
    sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
    return "\n".join([rows[0], sep] + rows[1:])


def extract_docx(path: str) -> str:
    """Walk the docx body in document order, emitting paragraphs and tables."""
    if _docx is None:
        raise RuntimeError("python-docx is not installed; cannot extract DOCX")

    doc = _docx.Document(path)
    body = doc.element.body

    chunks = []
    para_idx = 0
    table_idx = 0
    paragraphs = doc.paragraphs
    tables = doc.tables

    for child in body.iterchildren():
        tag = child.tag
        if tag == _qn("w:p"):
            if para_idx < len(paragraphs):
                text = paragraphs[para_idx].text.strip()
                if text:
                    chunks.append(text)
                para_idx += 1
        elif tag == _qn("w:tbl"):
            if table_idx < len(tables):
                md = _docx_table_to_markdown(tables[table_idx])
                if md:
                    chunks.append(md)
                table_idx += 1

    return "\n\n".join(chunks)


def extract_plain(path: str) -> str:
    """Read a plain UTF-8 text file (md, txt)."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def extract(path: str) -> Optional[str]:
    """
    Dispatch based on file extension.

    Returns:
        - extracted text for docx/md/txt
        - None for pdf (signals: agent will Read the original directly)
    """
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return None
    if ext == ".docx":
        return extract_docx(path)
    if ext in (".md", ".txt"):
        return extract_plain(path)
    raise ValueError(f"Unsupported file extension: {ext}")
